import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

brain_dir = r"C:\Users\Sys\.gemini\antigravity\brain\751ee7c4-1911-4d79-bb63-adfdecba8bcc"
output_docx = r"c:\Users\Sys\Desktop\my-startup-app\ShifaScribe_Week1_Report.docx"

doc = docx.Document()

# Page Setup - 1 inch margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Colors
COLOR_NAVY = RGBColor(15, 23, 42)      # #0F172A
COLOR_TEAL = RGBColor(13, 148, 136)    # #0D9488
COLOR_EMERALD = RGBColor(5, 150, 105)  # #059669
COLOR_GRAY = RGBColor(100, 116, 139)   # #64748B
COLOR_DARK = RGBColor(30, 41, 59)      # #1E293B
COLOR_LIGHT_BG = RGBColor(241, 245, 249) # #F1F5F9

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_TEAL
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p

def add_callout(text, bold_title=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    if bold_title:
        r_title = p.add_run(bold_title + " ")
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = COLOR_EMERALD
        
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10.5)
    r_text.font.italic = True
    r_text.font.color.rgb = COLOR_DARK
    return p

def add_image_figure(img_filename, caption_text, width_inches=5.8):
    path = os.path.join(brain_dir, img_filename)
    if os.path.exists(path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run(caption_text)
        r_cap.font.name = 'Calibri'
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = COLOR_GRAY

# ==================== 1. ENHANCED TITLE PAGE ====================
p_top_badge = doc.add_paragraph()
p_top_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_badge = p_top_badge.add_run("◆ MEDICAL AI ENGINEERING SPRINT • WEEK 1 COMPLETED ◆")
r_badge.font.name = 'Calibri'
r_badge.font.size = Pt(10)
r_badge.font.bold = True
r_badge.font.color.rgb = COLOR_EMERALD
p_top_badge.paragraph_format.space_before = Pt(10)
p_top_badge.paragraph_format.space_after = Pt(20)

# Main Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("ShifaScribe")
r_title.font.name = 'Calibri'
r_title.font.size = Pt(36)
r_title.font.bold = True
r_title.font.color.rgb = COLOR_NAVY
p_title.paragraph_format.space_after = Pt(4)

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("AI Urdu Voice-to-Text Medical Scribe System")
r_sub.font.name = 'Calibri'
r_sub.font.size = Pt(18)
r_sub.font.bold = True
r_sub.font.color.rgb = COLOR_TEAL
p_sub.paragraph_format.space_after = Pt(24)

# Decorative Line
p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
r_div.font.color.rgb = COLOR_EMERALD
r_div.font.bold = True

p_rep = doc.add_paragraph()
p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rep = p_rep.add_run("WEEK 1 TECHNICAL IMPLEMENTATION & PROGRESS REPORT")
r_rep.font.name = 'Calibri'
r_rep.font.size = Pt(13)
r_rep.font.bold = True
r_rep.font.color.rgb = COLOR_NAVY
p_rep.paragraph_format.space_before = Pt(18)
p_rep.paragraph_format.space_after = Pt(32)

# Formatted Metadata Card Table
t_card = doc.add_table(rows=6, cols=2)
t_card.alignment = WD_TABLE_ALIGNMENT.CENTER
t_card.autofit = False

card_data = [
    ("Prepared By (Lead Engineer):", "Abdul Rehman"),
    ("Submitted To (Supervisor):", "Khubaib Ahmed"),
    ("Project Name:", "ShifaScribe (Urdu Medical Speech AI)"),
    ("Reporting Scope:", "Week 1 Complete (Sprint 1: Days 1 to 5)"),
    ("Deployment Status:", "Active & Fully Verified"),
    ("Submission Date:", "August 12, 2026"),
]

for i, (label, val) in enumerate(card_data):
    row = t_card.rows[i]
    c1, c2 = row.cells[0], row.cells[1]
    c1.width = Inches(2.6)
    c2.width = Inches(3.7)
    
    if i % 2 == 0:
        set_cell_background(c1, "F1F5F9")
        set_cell_background(c2, "F1F5F9")
        
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(label)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10.5)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_NAVY
    
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(val)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10.5)
    r2.font.bold = (i in [0, 1])
    r2.font.color.rgb = COLOR_EMERALD if (i in [0, 1]) else COLOR_DARK

doc.add_page_break()

# ==================== 2. EXECUTIVE SUMMARY ====================
add_heading_1("1. Executive Summary")

add_body(
    "ShifaScribe is an AI-powered Urdu Voice-to-Text Medical Scribe System engineered specifically for high-throughput OPD (Outpatient Department) hospital clinics. In busy clinical environments, physicians spend up to 40% of consultation time manually typing clinical notes and symptoms into EHR systems. ShifaScribe eliminates this friction by listening to doctor-patient conversations spoken in natural Urdu, transcribing the speech in real-time, and structuring the data into standardized medical ICD-10 EHR clinical notes."
)

add_callout(
    "Sprint 1 Completed: All 5 Days of Week 1 technical milestones are 100% implemented, verified, and integrated. The full stack spans Next.js frontend UI, HTML5 MediaRecorder 16kHz audio capture, WebM compression, Python FastAPI backend, CORS middleware, PostgreSQL SQLAlchemy database schemas (patients, doctors, consultation_logs), and asynchronous audio file storage.",
    "Week 1 Summary:"
)

# ==================== 3. DAY 1 TECHNICAL IMPLEMENTATION ====================
add_heading_1("2. Day 1 Implementation: Workspace & Core UI Recorder")
add_body(
    "The application foundation was established using Next.js 16 App Router in combination with TypeScript and Tailwind CSS v4. The setup prioritizes instant page loading, minimal visual clutter for clinical users, and strict type safety across all components."
)

add_image_figure("code_recorder_snippet.png", "Figure 1: Code Implementation Screenshot of ConsultationRecorder.tsx (Day 1 Static UI Structure)")
add_image_figure("shifascribe_ui_idle_1786518121373.jpg", "Figure 2: ShifaScribe Doctor Consult Screen in Idle State (http://localhost:3000)")
add_image_figure("shifascribe_ui_rec_1786518136735.jpg", "Figure 3: Active OPD Consultation Recording State with Pulsing Waveform & Live Timer")
add_image_figure("shifascribe_ui_proc_1786518154916.jpg", "Figure 4: AI Urdu Speech Engine Processing State with EHR Output Preview")
add_image_figure("terminal_output.png", "Figure 5: Terminal Output Screenshot showing Local Server Initialization & Status 200 OK")

# ==================== 4. DAY 2 IMPLEMENTATION ====================
add_heading_1("3. Day 2 Implementation: Live Microphone Capture & MediaRecorder API")
add_body(
    "On Day 2, the static ConsultationRecorder component was upgraded to interface directly with browser audio hardware using navigator.mediaDevices.getUserMedia({ audio: true }). When the doctor clicks 'Start Consultation', the browser prompts for microphone permissions, captures the hardware audio stream, and initializes an HTML5 MediaRecorder instance."
)

add_image_figure("day2_code_snippet.png", "Figure 6: Day 2 Code Implementation of startRecording(), MediaRecorder events, and Blob URL generation")
add_image_figure("day2_console_output.png", "Figure 7: Browser Developer Console Output showing Permission Grant, Chunk Collection, and Blob Generation pipeline")

# ==================== 5. DAY 3 IMPLEMENTATION ====================
add_heading_1("4. Day 3 Implementation: 16kHz Mono Audio Compression & Git Sync")
add_body(
    "On Day 3, the audio capture pipeline was optimized to prevent network choke across congested hospital intranet links. By constraining getUserMedia with sampleRate: 16000, channelCount: 1 (Mono), echoCancellation: true, and noiseSuppression: true, raw microphone audio bandwidth is reduced by over 80% while retaining high speech intelligibility for speech-to-text engines."
)

add_image_figure("day3_code_snippet.png", "Figure 8: Day 3 Implementation of 16kHz Mono Constraints, WebM MIME Type Selection, and KB Size Logging")
add_image_figure("day3_compression_console.png", "Figure 9: Browser Developer Console Output verifying 16kHz Mono Constraints, Chunk Collection, KB Blob Size, and WebM MIME Type")

# ==================== 6. DAY 4 IMPLEMENTATION ====================
add_heading_1("5. Day 4 Implementation: Python FastAPI Backend & CORS Setup")
add_body(
    "On Day 4, the backend architecture was initialized inside backend/ using Python 3.14 and FastAPI. CORSMiddleware was configured to grant cross-origin requests from the Next.js frontend (allow_origins=['http://localhost:3000', 'http://127.0.0.1:3000']). A GET /health endpoint was implemented returning {'status': 'API is running', 'service': 'ShifaScribe Audio Engine', 'version': '0.1.0'}."
)

add_image_figure("day4_backend_main.png", "Figure 10: Day 4 FastAPI Boilerplate Implementation in backend/main.py showing CORSMiddleware & /health Route")
add_image_figure("day4_swagger_ui.png", "Figure 11: Uvicorn Server Execution & Terminal Verification of GET /health Endpoint (Status 200 OK)")
add_image_figure("day4_swagger_interactive_ui_1786520471196.jpg", "Figure 12: Interactive FastAPI Swagger OpenAPI UI Documentation (http://localhost:8000/docs) showing GET /health Response")
add_image_figure("day4_frontend_cors_ui_1786520497098.jpg", "Figure 13: Next.js Frontend Doctor Consult Screen UI (http://localhost:3000) displaying active FastAPI Backend Connection Badge")

# ==================== 7. DAY 5 IMPLEMENTATION ====================
add_heading_1("6. Day 5 Implementation: PostgreSQL Schemas & Async Audio Upload API")
add_body(
    "On Day 5, SQLAlchemy ORM models were constructed in backend/models.py to hold master clinical entities: Patient (ID, name, age, OPD token), Doctor (ID, name, department), and ConsultationLog (patient ID, doctor ID, audio file path, file size in KB, MIME format, and status). An asynchronous API endpoint POST /api/consultation/upload-audio was implemented saving audio files into storage/audio/."
)

add_image_figure("day5_models_code.png", "Figure 14: Day 5 SQLAlchemy ORM Model Schemas in backend/models.py (Patient, Doctor, ConsultationLog)")
add_image_figure("day5_upload_api.png", "Figure 15: Terminal Execution & Verification of POST /api/consultation/upload-audio returning HTTP 201 Created")
add_image_figure("day5_swagger_upload_ui_1786520971486.jpg", "Figure 16: Interactive Swagger OpenAPI Upload UI (http://localhost:8000/docs) displaying successful file upload payload")

# ==================== 8. CONCLUSION & NEXT STEPS ====================
add_heading_1("7. Week 1 Sprint Conclusion & Next Steps")
add_body(
    "Week 1 of ShifaScribe successfully established a complete full-stack clinical architecture. The frontend UI provides an intuitive, 3-state OPD doctor consultation interface with 16kHz mono audio compression. The FastAPI backend manages cross-origin communication, database schema persistence, and asynchronous binary audio file storage. Week 2 will focus on real-time OpenAI Whisper Urdu speech-to-text transcription streaming and ICD-10 medical entity extraction."
)

# ==================== 9. NEW SUPERVISOR REMARKS & APPROVAL SECTION ====================
add_heading_1("8. Supervisor Evaluation, Remarks & Approval")

add_body(
    "This section is reserved for the supervisor evaluation, comments, and formal sign-off for the Week 1 Technical Implementation of ShifaScribe."
)

# Formal Evaluation & Remarks Box Table
t_eval = doc.add_table(rows=5, cols=2)
t_eval.alignment = WD_TABLE_ALIGNMENT.CENTER
t_eval.autofit = False

# Header Row
hdr_cells = t_eval.rows[0].cells
hdr_cells[0].width = Inches(2.2)
hdr_cells[1].width = Inches(4.1)
set_cell_background(hdr_cells[0], "0F172A")
set_cell_background(hdr_cells[1], "0F172A")

p_h0 = hdr_cells[0].paragraphs[0]
r_h0 = p_h0.add_run("Evaluation Field")
r_h0.font.bold = True
r_h0.font.color.rgb = RGBColor(255, 255, 255)

p_h1 = hdr_cells[1].paragraphs[0]
r_h1 = p_h1.add_run("Supervisor Assessment & Details")
r_h1.font.bold = True
r_h1.font.color.rgb = RGBColor(255, 255, 255)

# Row 1: Rating
r1 = t_eval.rows[1]
r1.cells[0].width = Inches(2.2)
r1.cells[1].width = Inches(4.1)
p_r1_lbl = r1.cells[0].paragraphs[0].add_run("Overall Week 1 Rating:")
p_r1_lbl.font.bold = True
p_r1_val = r1.cells[1].paragraphs[0].add_run("[   ] Excellent    [   ] Very Good    [   ] Satisfactory    [   ] Needs Revision")

# Row 2: Remarks Box
r2 = t_eval.rows[2]
r2.cells[0].width = Inches(2.2)
r2.cells[1].width = Inches(4.1)
p_r2_lbl = r2.cells[0].paragraphs[0].add_run("Supervisor Remarks & Feedback:")
p_r2_lbl.font.bold = True

p_r2_val = r2.cells[1].paragraphs[0]
p_r2_val.paragraph_format.space_after = Pt(40)
r_rem = p_r2_val.add_run("(Please write supervisor remarks and evaluation notes here...)\n\n\n\n")
r_rem.font.italic = True
r_rem.font.color.rgb = COLOR_GRAY

# Row 3: Signature
r3 = t_eval.rows[3]
r3.cells[0].width = Inches(2.2)
r3.cells[1].width = Inches(4.1)
p_r3_lbl = r3.cells[0].paragraphs[0].add_run("Supervisor Signature:")
p_r3_lbl.font.bold = True
p_r3_val = r3.cells[1].paragraphs[0].add_run("_________________________________________\nKhubaib Ahmed")

# Row 4: Date & Approval
r4 = t_eval.rows[4]
r4.cells[0].width = Inches(2.2)
r4.cells[1].width = Inches(4.1)
p_r4_lbl = r4.cells[0].paragraphs[0].add_run("Date of Sign-off & Status:")
p_r4_lbl.font.bold = True
p_r4_val = r4.cells[1].paragraphs[0].add_run("Date: ______________, 2026   |   Status: [   ] APPROVED FOR WEEK 2")

# Save Document
doc.save(output_docx)
print("Successfully generated Word report with enhanced Title Page & Supervisor Remarks at:", output_docx)
