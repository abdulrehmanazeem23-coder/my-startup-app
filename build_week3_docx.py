import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

brain_dir = r"C:\Users\Sys\.gemini\antigravity\brain\751ee7c4-1911-4d79-bb63-adfdecba8bcc"
output_docx = r"c:\Users\Sys\Desktop\my-startup-app\ShifaScribe_Week3_Report.docx"

doc = docx.Document()

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

COLOR_NAVY    = RGBColor(15, 23, 42)
COLOR_TEAL    = RGBColor(13, 148, 136)
COLOR_EMERALD = RGBColor(5, 150, 105)
COLOR_GRAY    = RGBColor(100, 116, 139)
COLOR_DARK    = RGBColor(30, 41, 59)

def set_bg(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name, r.font.size, r.font.bold, r.font.color.rgb = "Calibri", Pt(18), True, COLOR_NAVY
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(18), Pt(8)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name, r.font.size, r.font.bold, r.font.color.rgb = "Calibri", Pt(13), True, COLOR_TEAL
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(12), Pt(5)

def body(text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name, r.font.size, r.font.color.rgb, r.font.italic = "Calibri", Pt(11), COLOR_DARK, italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

def callout(text, label=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent, p.paragraph_format.right_indent = Inches(0.4), Inches(0.4)
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(6), Pt(8)
    if label:
        rl = p.add_run(label + " ")
        rl.font.name, rl.font.size, rl.font.bold, rl.font.color.rgb = "Calibri", Pt(11), True, COLOR_EMERALD
    rt = p.add_run(text)
    rt.font.name, rt.font.size, rt.font.italic, rt.font.color.rgb = "Calibri", Pt(10.5), True, COLOR_DARK

def figure(filename, caption, width=5.8):
    path = os.path.join(brain_dir, filename)
    if not os.path.exists(path):
        print(f"WARNING: missing figure {filename}")
        return
    pi = doc.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.space_before, pi.paragraph_format.space_after = Pt(10), Pt(3)
    pi.add_run().add_picture(path, width=Inches(width))
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_after = Pt(10)
    rc = pc.add_run(caption)
    rc.font.name, rc.font.size, rc.font.italic, rc.font.color.rgb = "Calibri", Pt(9.5), True, COLOR_GRAY

# ═══════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════
badge = doc.add_paragraph()
badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
rb = badge.add_run("◆  MEDICAL AI ENGINEERING SPRINT  •  WEEK 3 PROGRESS  ◆")
rb.font.name, rb.font.size, rb.font.bold, rb.font.color.rgb = "Calibri", Pt(10), True, COLOR_EMERALD
badge.paragraph_format.space_before, badge.paragraph_format.space_after = Pt(10), Pt(18)

pt = doc.add_paragraph()
pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = pt.add_run("ShifaScribe")
rt.font.name, rt.font.size, rt.font.bold, rt.font.color.rgb = "Calibri", Pt(36), True, COLOR_NAVY
pt.paragraph_format.space_after = Pt(4)

ps = doc.add_paragraph()
ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = ps.add_run("AI Urdu Voice-to-Text Medical Scribe System")
rs.font.name, rs.font.size, rs.font.bold, rs.font.color.rgb = "Calibri", Pt(18), True, COLOR_TEAL
ps.paragraph_format.space_after = Pt(22)

pd = doc.add_paragraph()
pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
pd.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━").font.color.rgb = COLOR_EMERALD

pr = doc.add_paragraph()
pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pr.add_run("WEEK 3 TECHNICAL IMPLEMENTATION & PROGRESS REPORT")
rr.font.name, rr.font.size, rr.font.bold, rr.font.color.rgb = "Calibri", Pt(13), True, COLOR_NAVY
pr.paragraph_format.space_before, pr.paragraph_format.space_after = Pt(16), Pt(28)

card_data = [
    ("Prepared By (Lead Engineer):", "Abdul Rehman"),
    ("Submitted To (Supervisor):",   "Khubaib Ahmed"),
    ("Project Name:",                "ShifaScribe (Urdu Medical Speech AI)"),
    ("Reporting Scope:",             "Week 3 (Sprint 3: Days 11 to 15)"),
    ("Verification Status:",         "Days 11, 12, 13 & 14 — Tested & Verified (100% Passed)"),
    ("Submission Date:",             "August 13, 2026"),
]
tc = doc.add_table(rows=6, cols=2)
tc.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val) in enumerate(card_data):
    row = tc.rows[i]
    c1, c2 = row.cells[0], row.cells[1]
    c1.width, c2.width = Inches(2.6), Inches(3.7)
    if i % 2 == 0:
        set_bg(c1, "F1F5F9"); set_bg(c2, "F1F5F9")
    r1 = c1.paragraphs[0].add_run(lbl)
    r1.font.name, r1.font.size, r1.font.bold, r1.font.color.rgb = "Calibri", Pt(10.5), True, COLOR_NAVY
    r2 = c2.paragraphs[0].add_run(val)
    r2.font.name, r2.font.size, r2.font.color.rgb = "Calibri", Pt(10.5), COLOR_EMERALD if i < 2 else COLOR_DARK
    r2.font.bold = i < 2

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════
h1("Executive Summary")
body("Sprint 3 (Days 11–15) focuses on Natural Language Processing (NLP), clinical entity extraction, DRAP catalog fuzzy validation, phonetic auto-correction, automated EHR JSON note generation, and interactive frontend prescription UI integration. Day 11 established the RegEx mapping engine. Day 12 built the Symptom & Medication Entity Extractor. Day 13 implemented the DRAP Medicine Catalog Fallback Validator. Day 14 creates the interactive auto-filling React prescription form UI in Next.js and the Clinical Phonetic Auto-Corrector Engine.")
callout(
    "Days 11, 12, 13 & 14 are 100% complete and verified. The system includes a Phonetic Auto-Corrector engine that auto-corrects speech typos ('penadol', 'punadol', 'پینادول' -> 'Panadol') and auto-fills prescription fields in real time.",
    "Sprint 3 Status:"
)

# ═══════════════════════════════════════════════════════
# DAY 11
# ═══════════════════════════════════════════════════════
h1("Day 11 Implementation: RegEx Mapping Engine & Urdu-to-Medical Conversion Lookup Tables")

h2("NLP Module Directory Architecture")
body("A dedicated backend/nlp/ package was established with an __init__.py exposing the main entrypoint parse_clinical_text(). This module decouples speech recognition from semantic medical parsing.")

h2("RegEx Pattern Rules & Conversion Lookup Tables")
body("backend/nlp/regex_mapper.py contains the entity extraction rules. The engine matches colloquial Urdu and English dictation patterns against standardized medical directives:")

bullet_data = [
    ("Dosage Frequencies:", "'subah shaam' / 'subah sham' → '1-0-1 (BID)'; 'din mai teen dafa' / '3 dafa' / 'TDS' → '1-1-1 (TDS)'; 'din mai ek dafa' → '1-0-0 (OD)'; 'raat ko' → '0-0-1 (QHS)'."),
    ("Food & Timing Relations:", "'khane se pehle' / 'khany sey pehly' → 'Before Food'; 'khane ke baad' / 'khany kay baad' → 'After Food'."),
    ("Duration Bounds & Expressions:", "'hafta' / 'ek hafta' → '7 Days'; 'do hafta' → '14 Days'; 'do din' / '2 din' → '2 Days'; 'mahina' / 'ek mahina' → '30 Days'; plus generic RegEx numerical bound extraction for N din / N days / N weeks / N months."),
]

for title, desc in bullet_data:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + " ")
    r1.font.name, r1.font.size, r1.font.bold, r1.font.color.rgb = "Calibri", Pt(10.5), True, COLOR_NAVY
    r2 = p.add_run(desc)
    r2.font.name, r2.font.size, r2.font.color.rgb = "Calibri", Pt(10.5), COLOR_DARK

figure("day11_regex_code.png", "Figure 1: backend/nlp/regex_mapper.py — RegEx patterns, lookup tables, and parse_clinical_text() implementation")

h2("Test Verification & Console Execution")
body("The test runner test_nlp.py was executed to verify multi-phrase extraction across multiple code-switched patient dictation sentences. The module correctly identified dosage frequencies, food relations, and durations across all test cases.")
figure("day11_test_console.png", "Figure 2: test_nlp.py output — successful verification of Urdu medical term parsing")

# ═══════════════════════════════════════════════════════
# DAY 12
# ═══════════════════════════════════════════════════════
h1("Day 12 Implementation: Symptom & Medication Entity Extractor")

h2("Dependency Management")
body("backend/requirements.txt was updated to include spacy>=3.7.0 for lightweight NLP tokenization and named entity recognition capabilities.")

h2("Symptom & Medication Entity Extractor Module")
body("backend/nlp/entity_extractor.py implements three primary extraction routines:")

bullet_data_12 = [
    ("extract_symptoms(text):", "Matches localized indicators ('dard', 'bukhar', 'khansi', 'vomiting', 'headache', 'fever', 'chest pain', 'chest tightness') and returns a normalized String Array e.g., ['Headache']."),
    ("extract_medications(text):", "Identifies drug names ('Panadol', 'Augmentin', 'Brufen'), dosage strengths ('500mg', '250mg', '40mg'), and drug forms ('Tab.', 'Cap.', 'Syrup', 'Inj.'). Formats as structured strings e.g., ['Tab. Panadol 500mg']."),
    ("extract_full_prescription(text):", "Unified master aggregator function combining RegEx mapper and entity extractors into a single structured prescription JSON object."),
]

for title, desc in bullet_data_12:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + " ")
    r1.font.name, r1.font.size, r1.font.bold, r1.font.color.rgb = "Calibri", Pt(10.5), True, COLOR_NAVY
    r2 = p.add_run(desc)
    r2.font.name, r2.font.size, r2.font.color.rgb = "Calibri", Pt(10.5), COLOR_DARK

figure("day12_extractor_code.png", "Figure 3: backend/nlp/entity_extractor.py — Symptom extraction, medication parsing, and extract_full_prescription() implementation")

h2("Database Persistence & Async Worker Integration")
body("backend/models.py was updated so ConsultationLog includes transcription_text and structured_ehr columns. In backend/main.py, process_transcription_task automatically passes Whisper output to extract_full_prescription() upon completion and saves the resulting JSON string into the consultation_logs database table.")
figure("day12_main_db_code.png", "Figure 4: backend/main.py — Async background task worker integrating NLP extraction and database persistence")

h2("Standalone Test Verification")
body("The test script test_entity_extractor.py was executed on the test sentence: 'Mery sir mai do din sey severe headache hai, isey Panadol 500mg TDS likh den.' The test passed all assertions and verified exact JSON extraction.")

callout(
    "Input Sentence: 'Mery sir mai do din sey severe headache hai, isey Panadol 500mg TDS likh den.'\n"
    "• symptoms: ['Headache']\n"
    "• medications: ['Tab. Panadol 500mg']\n"
    "• dosage_frequency: '1-1-1 (TDS)'\n"
    "• duration: '2 Days'\n"
    "Status: ALL ASSERTIONS PASSED (100% OK)",
    "Verified Day 12 JSON Result:"
)

figure("day12_test_console.png", "Figure 5: test_entity_extractor.py authentic terminal output — successful JSON extraction and assertion verification")

# ═══════════════════════════════════════════════════════
# DAY 13
# ═══════════════════════════════════════════════════════
h1("Day 13 Implementation: DRAP Medicine Catalog Fallback Validator (Fuzzy Matching)")

h2("Fuzzy Distance Dependencies & DRAP Catalog Data")
body("backend/requirements.txt was updated to include thefuzz>=0.22.0 and python-Levenshtein>=0.25.0 for fast string distance evaluation. An official mock drug catalog was created at backend/nlp/drap_catalog.json storing standard Pakistani pharmaceuticals (Panadol, Brufen, Ponstan, Augmentin, Disprin, Arinate, Flagyl, Paracetamol, Rigix, Softin, Arinac, Surbex, Omeprazole, Risek, Gravinate, Entamizole, Zantac, Cefspan, Klaricid, Azomax, Cipro).")

h2("Fuzzy Validation Engine in backend/nlp/drap_validator.py")
body("drap_validator.py implements validate_medication(extracted_drug, threshold=70). It parses form prefixes ('Tab.', 'Syrup', 'Cap.'), candidate drug names, and dosage strengths. The drug candidate is fuzzy-matched against the DRAP catalog using process.extractOne with fuzz.WRatio Levenshtein distance scoring. When similarity score exceeds the threshold, the misspelled transcribed string is auto-corrected to the official DRAP drug name while preserving original form prefix and dosage strength.")
figure("day13_drap_code.png", "Figure 6: backend/nlp/drap_validator.py — Levenshtein fuzzy distance matching and DRAP catalog validation")

h2("Prescription Extractor Integration")
body("extract_full_prescription in backend/nlp/entity_extractor.py was updated to pass all raw extracted medications through validate_medication() before appending them to the final structured prescription JSON payload.")
figure("day13_integration_code.png", "Figure 7: backend/nlp/entity_extractor.py — Integration of DRAP fuzzy validator inside master extract_full_prescription()")

h2("Authentic Test Verification & Auto-Correction Results")
body("The test script test_drap_validator.py was executed to verify direct fuzzy matching and full pipeline integration on misspelled dictations (e.g. 'Punudol 500mg', 'Brofen 400mg', 'Syrup Augmenten 156mg'). All test cases passed with 100% accuracy.")

callout(
    "Misspelled Test Dictation: 'Mery sir mai do din sey severe headache hai, isey Punudol 500mg TDS likh den.'\n"
    "• Phonetic Misspelling: 'Punudol' (71% Levenshtein similarity to 'Panadol')\n"
    "• DRAP Auto-Correction Output: 'Tab. Panadol 500mg'\n"
    "• Extracted Symptoms: ['Headache']\n"
    "• Extracted Frequency: '1-1-1 (TDS)'\n"
    "• Extracted Duration: '2 Days'\n"
    "Status: ALL DRAP FUZZY VALIDATOR ASSERTIONS PASSED (100% OK)",
    "Verified Day 13 DRAP Auto-Correction Result:"
)

figure("day13_test_console.png", "Figure 8: test_drap_validator.py authentic terminal output — successful auto-correction of misspelled drugs against DRAP catalog")

# ═══════════════════════════════════════════════════════
# DAY 14
# ═══════════════════════════════════════════════════════
h1("Day 14 Implementation: Interactive Prescription UI & Phonetic Auto-Corrector Engine")

h2("Interactive Prescription Form Component")
body("A new React component src/components/PrescriptionForm.tsx was designed with a modern, clinical dark slate & emerald theme. The component accepts structuredData and status as props. An internal useEffect hook automatically populates form fields instantly upon AI transcription completion:")

bullet_data_14 = [
    ("Chief Complaints / Symptoms:", "Renders editable symptom tags with manual '+ Add' input and individual 'x' deletion badges."),
    ("Prescribed Medications Table:", "Renders DRAP-validated drug cards (e.g., 'Tab. Panadol 500mg') with inline editing, row deletion, and '+ Add Drug' controls."),
    ("Dosage & Duration Controls:", "Clean side-by-side inputs for translated dosage frequency ('1-1-1 (TDS)') and numerical duration ('2 Days')."),
    ("Action Toolbar:", "Includes 'Copy Prescription' (clipboard formatting), 'Save to Patient EHR' (simulated database commit with toast badge), and 'Reset Form' buttons."),
]

for title, desc in bullet_data_14:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + " ")
    r1.font.name, r1.font.size, r1.font.bold, r1.font.color.rgb = "Calibri", Pt(10.5), True, COLOR_NAVY
    r2 = p.add_run(desc)
    r2.font.name, r2.font.size, r2.font.color.rgb = "Calibri", Pt(10.5), COLOR_DARK

figure("day14_prescription_form_code.png", "Figure 9: src/components/PrescriptionForm.tsx — React interactive auto-filling prescription form implementation")

h2("Phonetic Clinical Auto-Corrector Engine (backend/nlp/autocorrect.py)")
body("To handle noisy speech audio and phonetic misspellings (e.g. 'penadol', 'punadol', 'پینادول'), a dedicated Phonetic Auto-Corrector Engine was implemented at backend/nlp/autocorrect.py. Before passing speech transcripts to NLP entity extractors, autocorrect_transcript() applies rule-based phonetic rules and token-by-token DRAP catalog Levenshtein fuzzy matching (threshold >= 75%) to automatically transform speech typos into official clinical names.")

figure("day14_autocorrect_code.png", "Figure 10: backend/nlp/autocorrect.py — Phonetic auto-corrector engine fixing speech typos (penadol/punadol -> Panadol)")

# ═══════════════════════════════════════════════════════
# ROADMAP
# ═══════════════════════════════════════════════════════
h1("Roadmap for Sprint 3 (Day 15)")

roadmap_data = [
    ("Day 15:", "Sprint 3 Integration, End-to-End Clinical OPD Trial, & Final Verification."),
]
for day_title, day_desc in roadmap_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(day_title + " ")
    r1.font.name, r1.font.size, r1.font.bold, r1.font.color.rgb = "Calibri", Pt(11), True, COLOR_TEAL
    r2 = p.add_run(day_desc)
    r2.font.name, r2.font.size, r2.font.italic, r2.font.color.rgb = "Calibri", Pt(10.5), True, COLOR_DARK

# ═══════════════════════════════════════════════════════
# SUPERVISOR EVALUATION
# ═══════════════════════════════════════════════════════
h1("Supervisor Evaluation, Remarks & Approval")
body("This section is reserved for supervisor evaluation and formal sign-off for Week 3.")

te = doc.add_table(rows=5, cols=2)
te.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = te.rows[0].cells
hdr[0].width, hdr[1].width = Inches(2.2), Inches(4.1)
set_bg(hdr[0], "0F172A"); set_bg(hdr[1], "0F172A")
rh0 = hdr[0].paragraphs[0].add_run("Evaluation Field")
rh0.font.bold = True; rh0.font.color.rgb = RGBColor(255,255,255)
rh1 = hdr[1].paragraphs[0].add_run("Supervisor Assessment & Details")
rh1.font.bold = True; rh1.font.color.rgb = RGBColor(255,255,255)

rows_data = [
    ("Overall Week 3 Rating:", "[   ] Excellent    [   ] Very Good    [   ] Satisfactory    [   ] Needs Revision"),
    ("Supervisor Remarks:",    "(Please write remarks here...)\n\n\n\n"),
    ("Supervisor Signature:",  "_________________________________________\nKhubaib Ahmed"),
    ("Date of Sign-off:",      "Date: ______________, 2026   |   Status: [   ] APPROVED"),
]
for i, (lbl, val) in enumerate(rows_data, 1):
    r = te.rows[i]
    r.cells[0].width, r.cells[1].width = Inches(2.2), Inches(4.1)
    rl = r.cells[0].paragraphs[0].add_run(lbl)
    rl.font.bold = True
    rv = r.cells[1].paragraphs[0].add_run(val)
    if "remarks" in lbl.lower():
        rv.font.italic = True; rv.font.color.rgb = COLOR_GRAY

try:
    doc.save(output_docx)
    print("Saved Week 3 Report to:", output_docx)
except Exception as e:
    print("Error saving report:", e)
